import re
import io
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors

def md_to_reportlab_html(text: str) -> str:
    """Converts basic markdown syntax to ReportLab HTML-like tags."""
    # Convert links [text](url) to <a href="url">text</a>
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2" color="#2563eb"><u>\1</u></a>', text)
    # Convert bold **text** to <b>text</b>
    text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
    # Convert italic *text* to <i>text</i>
    text = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', text)
    return text

def generate_pdf_from_md(md_text: str) -> io.BytesIO:
    """Generates a styled, print-ready PDF from markdown text."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=54,  # 0.75 inch
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Define a clean, premium color palette
    primary_color = colors.HexColor('#1e3a8a')  # Deep blue
    text_color = colors.HexColor('#334155')     # Slate gray
    title_color = colors.HexColor('#0f172a')    # Dark slate
    
    # Custom premium styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=title_color,
        alignment=TA_CENTER,
        spaceAfter=14
    )
    
    h1_style = ParagraphStyle(
        'DocH1',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=primary_color,
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True
    )
    
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=13,
        textColor=title_color,
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=text_color,
        spaceAfter=6
    )
    
    bullet_style = ParagraphStyle(
        'DocBullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=3
    )

    story = []
    lines = md_text.split('\n')
    
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            story.append(Spacer(1, 3))
            continue
            
        if line_strip.startswith('# '):
            title_text = md_to_reportlab_html(line_strip[2:])
            story.append(Paragraph(title_text, title_style))
        elif line_strip.startswith('## '):
            h1_text = md_to_reportlab_html(line_strip[3:])
            story.append(Paragraph(h1_text, h1_style))
        elif line_strip.startswith('### ') or line_strip.startswith('#### '):
            prefix_len = 4 if line_strip.startswith('### ') else 5
            h2_text = md_to_reportlab_html(line_strip[prefix_len:])
            story.append(Paragraph(h2_text, h2_style))
        elif line_strip.startswith('- ') or line_strip.startswith('* '):
            bullet_text = md_to_reportlab_html(line_strip[2:])
            story.append(Paragraph(f"&bull; {bullet_text}", bullet_style))
        else:
            body_text = md_to_reportlab_html(line_strip)
            story.append(Paragraph(body_text, body_style))
            
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_docx_from_md(md_text: str) -> io.BytesIO:
    """Generates a styled Word Document (.docx) from markdown text."""
    doc = Document()
    
    # Margin setup: 0.75 in on all sides
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    lines = md_text.split('\n')
    
    def parse_inline_runs(paragraph, text: str):
        # Match bold syntax **text**
        tokens = re.split(r'(\*\*[^*]+\*\*)', text)
        for token in tokens:
            if token.startswith('**') and token.endswith('**'):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            else:
                # Convert Markdown links [text](url) to just "text (url)" in Word
                clean_text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', token)
                paragraph.add_run(clean_text)

    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue
            
        if line_strip.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.alignment = 1  # Center
            run = p.add_run(line_strip[2:])
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(15, 23, 42)  # Dark slate
        elif line_strip.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line_strip[3:])
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(30, 58, 138)  # Deep blue
        elif line_strip.startswith('### ') or line_strip.startswith('#### '):
            prefix_len = 4 if line_strip.startswith('### ') else 5
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line_strip[prefix_len:])
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(15, 23, 42)
        elif line_strip.startswith('- ') or line_strip.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            parse_inline_runs(p, line_strip[2:])
            # Set font name to Arial
            for run in p.runs:
                run.font.name = 'Arial'
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            parse_inline_runs(p, line_strip)
            # Set font name to Arial
            for run in p.runs:
                run.font.name = 'Arial'
                
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
